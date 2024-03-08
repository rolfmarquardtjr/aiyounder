import streamlit as st
import pandas as pd
import docx
import fitz  # PyMuPDF
from openai import OpenAI

# Defina a função process_file antes de usá-la
def process_file(uploaded_file, file_type):
    """Processa o arquivo carregado e retorna seu texto."""
    text = ""
    if file_type == "PDF":
        with fitz.open(stream=uploaded_file.read()) as doc:
            text = " ".join(page.get_text() for page in doc)
    elif file_type == "Excel":
        df = pd.read_excel(uploaded_file)
        text = df.to_string(index=False)
    elif file_type == "Word":
        doc = docx.Document(uploaded_file)
        text = '\n'.join(para.text for para in doc.paragraphs)
    return text

# Solicita a chave da API da OpenAI através da interface do Streamlit
with st.sidebar:
    openai_api_key = st.text_input("Chave da API da OpenAI", key="chatbot_api_key", type="password")
    st.markdown("[Obtenha uma chave da API da OpenAI](https://platform.openai.com/account/api-keys)")
    st.markdown("[Veja o código fonte](https://github.com/streamlit/llm-examples/blob/main/Chatbot.py)")
    st.markdown("[![Abra no GitHub Codespaces](https://github.com/codespaces/badge.svg)](https://codespaces.new/streamlit/llm-examples?quickstart=1)")

    # UI para upload de documentos
    st.header("Upload de Documentos")
    file_types = ["PDF", "Excel", "Word"]
    file_type = st.selectbox("Tipo de Documento", ["Escolher"] + file_types)
    uploaded_file = st.file_uploader("Carregue um arquivo", type=["pdf", "xlsx", "docx"])
    if uploaded_file and file_type != "Escolher":
        document_content = process_file(uploaded_file, file_type)
        st.session_state['document_content'] = document_content

st.title("💬 Chatbot")

st.title("💬 Chatbot")

# Função para processar o arquivo carregado
def process_file(uploaded_file, file_type):
    """Processa o arquivo carregado e retorna seu texto."""
    text = ""
    if file_type == "PDF":
        with fitz.open(stream=uploaded_file.read()) as doc:
            text = " ".join(page.get_text() for page in doc)
    elif file_type == "Excel":
        df = pd.read_excel(uploaded_file)
        text = df.to_string(index=False)
    elif file_type == "Word":
        doc = docx.Document(uploaded_file)
        text = '\n'.join(para.text for para in doc.paragraphs)
    return text

# Inicialização do estado da sessão para mensagens
if "messages" not in st.session_state:
    st.session_state["messages"] = [{"role": "assistant", "content": "Como posso ajudá-lo?"}]

# Exibição das mensagens anteriores
for msg in st.session_state.messages:
    st.chat_message(msg["role"]).write(msg["content"])

# Entrada de texto do usuário
if prompt := st.chat_input("Digite sua pergunta ou peça para resumir o documento carregado:"):
    if not openai_api_key:
        st.info("Por favor, adicione sua chave da API da OpenAI para continuar.")
        st.stop()

    # Se o usuário pedir para resumir o documento, use o conteúdo do documento como entrada
    if "resumir" in prompt.lower() and 'document_content' in st.session_state:
        prompt = st.session_state['document_content']

    # Adiciona a pergunta do usuário ao estado da sessão
    st.session_state.messages.append({"role": "user", "content": prompt})
    st.chat_message("user").write(prompt)

    # Criação do cliente da OpenAI com a chave fornecida
    client = OpenAI(api_key=openai_api_key)
    
    # Chama a API da OpenAI para obter a resposta
    response = client.chat.completions.create(model="gpt-3.5-turbo", messages=st.session_state.messages)
    msg = response.choices[0].message.content
    
    # Adiciona a resposta do assistente ao estado da sessão
    st.session_state.messages.append({"role": "assistant", "content": msg})
    st.chat_message("assistant").write(msg)
